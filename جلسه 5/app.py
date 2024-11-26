from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re

# تنظیمات Flask
app = Flask(__name__)

# تابع اصلاح نام فایل (حذف کاراکترهای غیرمجاز)
def اصلاح_نام_فايل(نام_فايل):
    return re.sub(r'[\\/:"*?<>|]', '-', نام_فايل)  # جایگزینی کاراکترهای غیرمجاز با "-"

# تابع ایجاد فایل Word
def ايجاد_فايل_ورد(data, مسير_ذخيره_سازي, نام_فايل):
    doc = Document()

    # بسم الله الرحمن الرحیم وسط و بزرگتر
    عنوان = doc.add_paragraph()
    run_عنوان = عنوان.add_run('بسم الله الرحمن الرحیم')
    run_عنوان.font.size = Pt(30)  # تنظیم سایز فونت
    run_عنوان.bold = True  # بولد کردن
    عنوان.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # وسط چین

    # ایجاد جدول برای داده‌ها
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # راست‌چین کردن جدول

    # اضافه کردن header برای جدول
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'مقدار'
    hdr_cells[1].text = 'مورد'

    # اضافه کردن داده‌ها به جدول
    for key, value in data.items():
        if value:
            row = table.add_row().cells
            row[0].text = value
            row[1].text = key
            for cell in row:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(14)  # سایز فونت سلول‌ها

    # ذخیره فایل Word
    file_path = os.path.join(مسير_ذخيره_سازي, f'{نام_فايل}.docx')
    doc.save(file_path)
    return file_path

# صفحه اصلی فرم
@app.route('/')
def index():
    return render_template('index.html')

# دریافت اطلاعات از فرم و ایجاد فایل Word
@app.route('/submit', methods=['POST'])
def submit():
    # اطمینان از اینکه فرم به درستی ارسال شده باشد
    if 'filename' not in request.form or 'date' not in request.form or 'this_week_report' not in request.form:
        return "خطا: همه فیلدهای فرم باید پر شوند."

    attendees = request.form.getlist('attendees')  # دریافت افراد حاضر از چک‌باکس‌ها
    all_members = ['امیر محمد', 'محمد فواد', 'طاها', 'امیر علی']
    absentee = [name for name in all_members if name not in attendees]  # محاسبه افراد غایب

    data = {
        'تاریخ جلسه': request.form['date'],
        'افراد حاضر': ', '.join(attendees),
        'افراد غایب': ', '.join(absentee),
        'گزارش این هفته': request.form['this_week_report'],
        'گزارش هفته آینده': request.form['next_week_report'],
        'پیگیری هفته آینده': request.form['follow_up']
    }

    مسير_ذخيره_سازي = os.path.join(os.getcwd(), 'generated_files')  # مسیر ذخیره‌سازی فایل‌ها
    os.makedirs(مسير_ذخيره_سازي, exist_ok=True)  # اگر پوشه وجود ندارد، آن را ایجاد کن
    نام_فايل = اصلاح_نام_فايل(request.form['filename'])  # اصلاح نام فایل

    if نام_فايل:
        file_path = ايجاد_فايل_ورد(data, مسير_ذخيره_سازي, نام_فايل)
        return send_file(file_path, as_attachment=True)
    else:
        return "خطا: نام فایل را وارد کنید."

if __name__ == '__main__':
    app.run(debug=True)
